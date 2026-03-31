import os
import re
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordGenerator:
    def __init__(self, output_dir="output_docs"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def create_summary_word(self, structured_summary: str, codigo_rastreio: str = None, username: str = "", corpo_resposta: str = "", meta: dict = None) -> str:
        os.makedirs(self.output_dir, exist_ok=True)
        nome_arquivo = f"esboco_{codigo_rastreio}.docx" if codigo_rastreio else f"{uuid.uuid4().hex}.docx"
        caminho_docx = os.path.join(self.output_dir, nome_arquivo)
        
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        if meta is None:
            meta = {}
        
        ano_atual = datetime.now().year
        num_memo = codigo_rastreio[:4].upper() if codigo_rastreio else "0000"
        
        p_memo = doc.add_paragraph()
        p_memo.add_run(f"Memo n.º {num_memo}/{ano_atual}/SETI-")
        p_memo.add_run("[INSIRA AQUI]").bold = True
        p_memo.paragraph_format.space_after = Pt(12)
        
        p_para = doc.add_paragraph()
        p_para.add_run("Para: ")
        p_para.add_run("[INSIRA AQUI]").bold = True
        p_para.paragraph_format.space_after = Pt(12)
        
        assunto_bruto = meta.get('Assunto', '').strip()
        assunto_limpo = re.sub(r'[:_\-\.]+$', '', assunto_bruto).strip()
        palavras_ignoradas = ["n/a", "detalhamento", "assunto", ""]
        if not assunto_limpo or assunto_limpo.lower() in palavras_ignoradas:
            assunto_final = 'Documento Oficial'
        else:
            assunto_final = assunto_limpo
            
        p_assunto = doc.add_paragraph(f"Assunto: {assunto_final}")
        p_assunto.paragraph_format.space_after = Pt(24)
        
        p_saudacao = doc.add_paragraph("Senhor(a),")
        p_saudacao.paragraph_format.space_after = Pt(12)
        
        if corpo_resposta:
            corpo_formatado = corpo_resposta.replace('[INSERIR AQUI]', '[INSIRA AQUI]')
            linhas_corpo = corpo_formatado.split('\n')
            for linha in linhas_corpo:
                if linha.strip():
                    p_corpo = doc.add_paragraph()
                    p_corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_corpo.paragraph_format.first_line_indent = Inches(0.5)
                    p_corpo.paragraph_format.space_after = Pt(12)
                    
                    partes = linha.split('[INSIRA AQUI]')
                    for idx, parte in enumerate(partes):
                        p_corpo.add_run(parte)
                        if idx < len(partes) - 1:
                            p_corpo.add_run('[INSIRA AQUI]').bold = True

        p_espaco = doc.add_paragraph()
        p_espaco.paragraph_format.space_after = Pt(24)
        
        meses = ["", "janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
        hoje = datetime.now()
        data_str = f"Curitiba, {hoje.day} de {meses[hoje.month]} de {hoje.year}."
        
        p_data = doc.add_paragraph(data_str)
        p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_data.paragraph_format.space_after = Pt(48)
        
        p_user = doc.add_paragraph()
        p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_user = p_user.add_run(username.upper() if username else "USUÁRIO")
        run_user.bold = True
        p_user.paragraph_format.space_after = Pt(2)
        
        p_setor = doc.add_paragraph()
        p_setor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_setor.add_run("[INSIRA AQUI]").bold = True
        p_setor.add_run("/Seti")
        
        doc.save(caminho_docx)
        return caminho_docx